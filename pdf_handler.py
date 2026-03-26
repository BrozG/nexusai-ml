"""
PDF/DOCX/TXT Handler for Document Processing
- Extracts text from documents and saves to raw_data/
- Optionally saves original files to policies/
"""

import logging
import re
import sys
from pathlib import Path
from typing import Dict, Optional, Union, BinaryIO

# Set UTF-8 for Windows
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Optional imports with availability flags
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocumentHandler:
    """Handles document processing and text extraction."""

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt'}

    def __init__(self, base_dir: str = "."):
        self.base_dir = Path(base_dir)
        self.raw_data_dir = self.base_dir / "raw_data"
        self.policies_dir = self.base_dir / "policies"

    def _ensure_directory(self, base_path: Path, domain: str, company: str) -> Path:
        """Create necessary directory structure."""
        output_path = base_path / domain / company
        output_path.mkdir(parents=True, exist_ok=True)
        return output_path

    def _clean_filename(self, filename: str) -> str:
        """Extract clean filename without extension."""
        name = Path(filename).stem
        name = re.sub(r'[^\w\-_]', '_', name)
        name = re.sub(r'_+', '_', name)
        name = name[:100]
        return name if name else 'document'

    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        cleaned = '\n'.join(lines)
        cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
        cleaned = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', cleaned)
        return cleaned.strip()

    def _extract_pdf(self, file_content: bytes) -> Optional[str]:
        """Extract text from PDF file."""
        if not PDF_AVAILABLE:
            logger.error("PyPDF2 not installed. Run: pip install PyPDF2")
            return None

        try:
            import io
            pdf_file = io.BytesIO(file_content)
            reader = PyPDF2.PdfReader(pdf_file)

            text_parts = []
            for page_num, page in enumerate(reader.pages, 1):
                try:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                except Exception as e:
                    logger.warning(f"Failed to extract page {page_num}: {e}")
                    continue

            full_text = '\n'.join(text_parts)
            logger.info(f"Extracted {len(full_text)} characters from {len(reader.pages)} pages")
            return full_text

        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return None

    def _extract_docx(self, file_content: bytes) -> Optional[str]:
        """Extract text from DOCX file."""
        if not DOCX_AVAILABLE:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return None

        try:
            import io
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)

            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            full_text = '\n'.join(text_parts)
            logger.info(f"Extracted {len(full_text)} characters from DOCX")
            return full_text

        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return None

    def _extract_txt(self, file_content: bytes) -> Optional[str]:
        """Extract text from TXT file."""
        try:
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                text = file_content.decode('latin-1', errors='ignore')
            logger.info(f"Extracted {len(text)} characters from TXT")
            return text
        except Exception as e:
            logger.error(f"TXT extraction failed: {e}")
            return None

    def extract_text(self, file_content: bytes, filename: str, domain: str, company: str) -> Dict:
        """Extract text from document and save to raw_data/"""
        try:
            file_ext = Path(filename).suffix.lower()
            if file_ext not in self.SUPPORTED_EXTENSIONS:
                return {
                    "success": False,
                    "message": f"Unsupported: {file_ext}. Use: {', '.join(self.SUPPORTED_EXTENSIONS)}",
                    "output_path": None
                }

            # Extract text
            if file_ext == '.pdf':
                text = self._extract_pdf(file_content)
            elif file_ext == '.docx':
                text = self._extract_docx(file_content)
            elif file_ext == '.txt':
                text = self._extract_txt(file_content)
            else:
                return {"success": False, "message": f"Unknown type: {file_ext}", "output_path": None}

            if text is None:
                return {"success": False, "message": f"Failed to extract from {filename}", "output_path": None}

            cleaned_text = self._clean_text(text)
            if not cleaned_text:
                return {"success": False, "message": "No text content found", "output_path": None}

            # Save extracted text
            output_dir = self._ensure_directory(self.raw_data_dir, domain, company)
            clean_name = self._clean_filename(filename)
            output_file = output_dir / f"pdf_{clean_name}.txt"

            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(cleaned_text)

            logger.info(f"Saved {len(cleaned_text)} characters to {output_file}")

            return {
                "success": True,
                "message": f"Processed {filename}",
                "output_path": str(output_file.relative_to(self.base_dir)),
                "characters_extracted": len(cleaned_text)
            }

        except Exception as e:
            logger.error(f"Error processing {filename}: {e}")
            return {"success": False, "message": f"Error: {str(e)}", "output_path": None}

    def save_original(self, file_content: bytes, filename: str, domain: str, company: str) -> Dict:
        """Save original file to policies/"""
        try:
            policies_dir = self._ensure_directory(self.policies_dir, domain, company)
            original_path = policies_dir / filename

            with open(original_path, 'wb') as f:
                f.write(file_content)

            logger.info(f"Original saved: {original_path}")
            return {
                "success": True,
                "original_path": str(original_path.relative_to(self.base_dir))
            }
        except Exception as e:
            logger.error(f"Failed to save original: {e}")
            return {"success": False, "original_path": None}


# === Public API Functions ===

def handle_pdf(
    file: Union[bytes, BinaryIO],
    domain: str,
    company: str,
    filename: Optional[str] = None
) -> Dict:
    """
    Extract text from PDF/DOCX/TXT (saves to raw_data/ only).

    Args:
        file: File bytes or file-like object
        domain: Domain category
        company: Company name
        filename: Original filename

    Returns:
        Dict with success, message, output_path, characters_extracted
    """
    handler = DocumentHandler()

    if hasattr(file, 'read'):
        file_content = file.read()
        if hasattr(file, 'seek'):
            file.seek(0)
        if filename is None and hasattr(file, 'filename'):
            filename = file.filename
    else:
        file_content = file

    if filename is None:
        return {"success": False, "message": "Filename required", "output_path": None}

    return handler.extract_text(file_content, filename, domain, company)


def handle_pdf_with_original(
    file: Union[bytes, BinaryIO],
    domain: str,
    company: str,
    filename: str,
    base_dir: str = "."
) -> Dict:
    """
    Save original file + extract text.

    Saves to:
    - policies/{domain}/{company}/{filename} (original)
    - raw_data/{domain}/{company}/pdf_{name}.txt (extracted)

    Args:
        file: File bytes or file-like object
        domain: Domain category
        company: Company name
        filename: Original filename
        base_dir: Base directory

    Returns:
        Dict with original_path, extracted_path, characters_extracted
    """
    handler = DocumentHandler(base_dir)

    if hasattr(file, 'read'):
        file_content = file.read()
        if hasattr(file, 'seek'):
            file.seek(0)
    else:
        file_content = file

    # Save original
    original_result = handler.save_original(file_content, filename, domain, company)

    # Extract text
    extract_result = handler.extract_text(file_content, filename, domain, company)

    return {
        "success": extract_result["success"] and original_result["success"],
        "message": extract_result["message"],
        "original_path": original_result.get("original_path"),
        "extracted_path": extract_result.get("output_path"),
        "characters_extracted": extract_result.get("characters_extracted"),
        "domain": domain,
        "company": company,
        "filename": filename
    }


def list_pdfs(domain: str, company: str, base_dir: str = ".") -> Dict:
    """List all original files in policies/{domain}/{company}/"""
    policies_dir = Path(base_dir) / "policies" / domain / company

    if not policies_dir.exists():
        return {"domain": domain, "company": company, "count": 0, "files": []}

    files = []
    for f in sorted(policies_dir.glob("*.*")):
        if f.is_file():
            stat = f.stat()
            files.append({
                "filename": f.name,
                "size_mb": round(stat.st_size / 1024 / 1024, 2),
                "extension": f.suffix
            })

    return {"domain": domain, "company": company, "count": len(files), "files": files}


def get_original_path(domain: str, company: str, filename: str, base_dir: str = ".") -> Path:
    """Get path to original file in policies/"""
    return Path(base_dir) / "policies" / domain / company / filename


# === CLI ===

def main():
    import argparse

    parser = argparse.ArgumentParser(
        description='Document Handler - Process PDF/DOCX/TXT files',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Extract text only
  python pdf_handler.py --file doc.pdf --domain ecommerce --company amazon

  # Save original + extract text
  python pdf_handler.py --file doc.pdf --domain ecommerce --company amazon --save-original

  # List saved files
  python pdf_handler.py --list --domain ecommerce --company amazon
        """
    )

    parser.add_argument('--file', help='Path to input file')
    parser.add_argument('--domain', help='Domain category')
    parser.add_argument('--company', help='Company name')
    parser.add_argument('--save-original', action='store_true', help='Also save original file to policies/')
    parser.add_argument('--list', action='store_true', help='List saved PDFs for domain/company')

    args = parser.parse_args()

    if args.list:
        if not args.domain or not args.company:
            print("Error: --list requires --domain and --company")
            sys.exit(1)
        result = list_pdfs(args.domain, args.company)
        print(f"\nFiles for {args.domain}/{args.company}: {result['count']}\n")
        for f in result['files']:
            print(f"  {f['filename']} ({f['size_mb']} MB)")
    else:
        if not args.file or not args.domain or not args.company:
            parser.print_help()
            sys.exit(1)

        file_path = Path(args.file)
        if not file_path.exists():
            print(f"Error: File not found: {args.file}")
            sys.exit(1)

        with open(file_path, 'rb') as f:
            if args.save_original:
                result = handle_pdf_with_original(f, args.domain, args.company, file_path.name)
                if result['success']:
                    print(f"\nOriginal: {result['original_path']}")
                    print(f"Extracted: {result['extracted_path']}")
                    print(f"Characters: {result['characters_extracted']}")
                else:
                    print(f"Error: {result['message']}")
                    sys.exit(1)
            else:
                result = handle_pdf(f, args.domain, args.company, file_path.name)
                if result['success']:
                    print(f"\nExtracted: {result['output_path']}")
                    print(f"Characters: {result['characters_extracted']}")
                else:
                    print(f"Error: {result['message']}")
                    sys.exit(1)


if __name__ == "__main__":
    main()
